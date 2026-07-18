import docx

def read_docx(filepath):
    doc = docx.Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # 读取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            full_text.append(" | ".join(row_text))
    
    return '\n'.join(full_text)

print("="*80)
print("PODOPENER.docx 完整内容")
print("="*80)
print(read_docx("PODOPENER.docx"))

print("\n" + "="*80)
print("alarmnew.docx 完整内容")
print("="*80)
print(read_docx("alarmnew.docx"))