import docx

def read_docx_tables(filepath):
    doc = docx.Document(filepath)
    
    print("="*80)
    print(f"文件: {filepath}")
    print("="*80)
    
    print("\n【段落内容】")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"{i}: {para.text}")
    
    print("\n【表格内容】")
    for t_idx, table in enumerate(doc.tables):
        print(f"\n表格 {t_idx + 1}:")
        for r_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            print(f"  行{r_idx}: {' | '.join(row_data)}")

read_docx_tables("PODOPENER.docx")